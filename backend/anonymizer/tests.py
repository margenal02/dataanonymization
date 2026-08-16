import io
import json
import shutil
import tempfile
import zipfile
from types import ModuleType
from pathlib import Path
from unittest.mock import patch

import xlrd
import xlwt
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

from .file_processors import ProcessingError, process_file
from .crypto import decrypt_mapping, encrypt_mapping
from .models import AnonymizationTask, RecognitionLabel, TrainingDocument, TrainingExample
from .paddlenlp_compat import ensure_aistudio_download_compatibility
from .ppstructure_worker import _build_pipeline as build_ppstructure_pipeline
from .ppstructure_worker import _extract_text as extract_ppstructure_text
from .ppstructure_worker import _recognize as recognize_with_ppstructure
from .recognizer import MappingBuilder, restore_text, suggest_organization_alias_groups
from .training_data import decrypt_label
from .uie_runtime import UIEProcessingError, _manager_url
from .uie_worker import _predict
from .views import _select_model_entities


class MappingBuilderTests(TestCase):
    @override_settings(MAPPING_ENCRYPTION_KEY="machine-a-secret")
    def test_encrypted_mapping_cannot_be_opened_with_another_machine_key(self):
        ciphertext = encrypt_mapping({"token_to_original": {"【LOCAL-人001】": "张三"}})
        with override_settings(MAPPING_ENCRYPTION_KEY="machine-b-secret"):
            with self.assertRaisesRegex(ValueError, "部署密钥"):
                decrypt_mapping(ciphertext)

    def test_detects_tobacco_entities_and_restores_exactly(self):
        source = "单位：中国烟草总公司\n联系人：张三\n电话：13800138000"
        builder = MappingBuilder("abcd1234")
        anonymized = builder.anonymize(source)
        self.assertNotIn("中国烟草总公司", anonymized)
        self.assertNotIn("张三", anonymized)
        self.assertNotIn("13800138000", anonymized)
        self.assertIn("【单001】", anonymized)
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_detects_entities_in_real_tobacco_document_language(self):
        source = (
            "云南中烟工业有限责任公司与红云红河烟草（集团）有限责任公司签订采购合同，"
            "项目负责人王建国，李娜负责复核，送货至云南省昆明市五华区红锦路123号，"
            "联系电话0871-63568888。"
        )
        builder = MappingBuilder("real1234")
        anonymized = builder.anonymize(source)

        for sensitive_value in (
            "云南中烟工业有限责任公司",
            "红云红河烟草（集团）有限责任公司",
            "王建国",
            "李娜",
            "云南省昆明市五华区红锦路123号",
            "0871-63568888",
        ):
            self.assertNotIn(sensitive_value, anonymized)
        self.assertEqual(builder.counts(), {"单位": 2, "人名": 2, "地址": 1, "电话": 1})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_detects_names_on_standalone_table_like_lines_and_name_lists(self):
        source = "人员名单：张三、李四、王五\n\n欧阳娜\n部门：财务部"
        builder = MappingBuilder("table123")
        anonymized = builder.anonymize(source)

        for sensitive_value in ("张三", "李四", "王五", "欧阳娜", "财务部"):
            self.assertNotIn(sensitive_value, anonymized)
        self.assertEqual(builder.counts(), {"人名": 4, "单位": 1})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_detects_group_leader_and_member_names_without_masking_role_phrases(self):
        source = (
            "组长：潘富昆。负责整个技能竞赛的宏观指导与活动协调。\n"
            "成员：赵英桥、李作英。负责跟进各项活动的具体实施。\n"
            "参加活动人员：厂领导、【单002】及生产加工车间参赛人员。"
        )
        builder = MappingBuilder("fca2002", ["person"])
        anonymized = builder.anonymize(source)

        for sensitive_value in ("潘富昆", "赵英桥", "李作英"):
            self.assertNotIn(sensitive_value, anonymized)
        self.assertIn("厂领导", anonymized)
        self.assertIn("生产加工车间参赛人员", anonymized)
        self.assertEqual(builder.counts(), {"人名": 3})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_anonymizes_person_in_filename_stem_and_restores_with_same_mapping(self):
        builder = MappingBuilder("file1234", ["person"])
        anonymized_stem = builder.anonymize_filename_stem("潘富昆技能竞赛活动方案")

        self.assertNotIn("潘富昆", anonymized_stem)
        self.assertIn("ANON_人001", anonymized_stem)
        self.assertEqual(restore_text(anonymized_stem, builder.export()), "潘富昆技能竞赛活动方案")

    def test_registers_valid_model_entity_and_rejects_role_word(self):
        builder = MappingBuilder("model123", ["person"])
        self.assertIsNotNone(builder.register_detected("潘富昆", "person"))
        self.assertIsNone(builder.register_detected("工作人员", "person"))
        self.assertIsNone(builder.register_detected("文山雨露", "person"))
        self.assertEqual(builder.counts(), {"人名": 1})

    def test_reviewed_aliases_share_one_token_and_restore_to_canonical_name(self):
        builder = MappingBuilder("alias123", ["organization"])
        builder.register("陆良复烤厂", "organization")
        builder.register("陆良厂", "organization")
        token = builder.merge_aliases("陆良复烤厂", ["陆良复烤厂", "陆良厂"])

        anonymized = builder.anonymize("陆良复烤厂安排陆良厂复核。")
        self.assertEqual(anonymized.count(token), 2)
        self.assertEqual(builder.counts(), {"单位": 1})
        self.assertEqual(restore_text(anonymized, builder.export()), "陆良复烤厂安排陆良复烤厂复核。")

    def test_recognizes_industry_abbreviations_without_full_names(self):
        source = "陆良厂与云南复烤、贵州复烤、河北中烟联合开展技术研究。"
        builder = MappingBuilder("industry-alias", ["organization"])

        anonymized = builder.anonymize(source)

        for abbreviation in ("陆良厂", "云南复烤", "贵州复烤", "河北中烟"):
            self.assertNotIn(abbreviation, anonymized)
            self.assertIn(abbreviation, builder.original_to_token)
        self.assertEqual(builder.counts(), {"单位": 4})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_suggests_rebake_and_china_tobacco_full_name_aliases(self):
        source = (
            "云南烟叶复烤有限责任公司简称云南复烤；"
            "贵州烟叶复烤有限责任公司简称贵州复烤；"
            "河北中烟工业有限责任公司简称河北中烟；"
            "陆良复烤厂后文简称陆良厂。"
        )
        builder = MappingBuilder("industry-groups", ["organization"])
        builder.discover(source)

        groups = suggest_organization_alias_groups(builder, source)
        pairs = {(group["canonical"], group["members"][1]) for group in groups}

        self.assertTrue({
            ("云南烟叶复烤有限责任公司", "云南复烤"),
            ("贵州烟叶复烤有限责任公司", "贵州复烤"),
            ("河北中烟工业有限责任公司", "河北中烟"),
            ("陆良复烤厂", "陆良厂"),
        }.issubset(pairs))

    def test_recognizes_factory_names_with_one_shared_trailing_suffix(self):
        source = "陆良、泸西、石林厂联合开展技术改造。"
        builder = MappingBuilder("shared-factory", ["organization"])

        anonymized = builder.anonymize(source)

        for abbreviation in ("陆良", "泸西", "石林"):
            self.assertIn(abbreviation, builder.original_to_token)
            self.assertNotIn(abbreviation, anonymized)
        self.assertIn("、", anonymized)
        self.assertIn("厂联合开展技术改造", anonymized)
        self.assertEqual(builder.counts(), {"单位": 3})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_preserves_factory_count_and_colloquial_grammar_while_masking_names(self):
        source = "公司先后走访陆良、泸西、石林等3家复烤厂，毕节、曲靖两厂同时参与。"
        builder = MappingBuilder("factory-count", ["organization"])

        anonymized = builder.anonymize(source)

        for abbreviation in ("陆良", "泸西", "石林", "毕节", "曲靖"):
            self.assertIn(abbreviation, builder.original_to_token)
            self.assertNotIn(abbreviation, anonymized)
        self.assertIn("等3家复烤厂", anonymized)
        self.assertIn("两厂同时参与", anonymized)
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_suggests_full_names_for_contextual_shared_suffix_abbreviations(self):
        source = (
            "陆良复烤厂、泸西复烤厂、石林复烤厂完成改造，"
            "后文统称陆良、泸西、石林三厂。"
        )
        builder = MappingBuilder("shared-groups", ["organization"])
        builder.discover(source)

        groups = suggest_organization_alias_groups(builder, source)
        pairs = {(group["canonical"], group["members"][1]) for group in groups}

        self.assertTrue({
            ("陆良复烤厂", "陆良"),
            ("泸西复烤厂", "泸西"),
            ("石林复烤厂", "石林"),
        }.issubset(pairs))
        for group in groups:
            if group["members"][1] in {"陆良", "泸西", "石林"}:
                self.assertIn("顿号并列", group["reason"])

    def test_explicit_factory_list_is_not_misread_as_suffix_omission(self):
        source = "陆良复烤厂、泸西复烤厂、石林复烤厂联合开展工作。"
        builder = MappingBuilder("full-factory-list", ["organization"])

        anonymized = builder.anonymize(source)

        self.assertEqual(
            set(builder.original_to_token),
            {"陆良复烤厂", "泸西复烤厂", "石林复烤厂"},
        )
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_cleans_sentence_prefix_before_rebake_company_name(self):
        source = "遵义复烤厂中心实验室为贵州烟叶复烤有限责任公司中心实验室。"
        builder = MappingBuilder("clean-org", ["organization"])

        builder.discover(source)

        self.assertIn("贵州烟叶复烤有限责任公司", builder.original_to_token)
        self.assertNotIn("中心实验室为贵州烟叶复烤有限责任公司", builder.original_to_token)

    def test_alias_merge_does_not_reuse_an_active_token_number(self):
        builder = MappingBuilder("alias-sequence", ["organization"])
        for name in ("甲单位", "乙单位", "丙单位"):
            builder.register(name, "organization")
        third_token = builder.original_to_token["丙单位"]

        builder.merge_aliases("甲单位", ["甲单位", "乙单位"])
        fourth_token = builder.register("丁单位", "organization")

        self.assertNotEqual(fourth_token, third_token)
        self.assertEqual(builder.token_to_original[third_token], "丙单位")
        self.assertEqual(builder.token_to_original[fourth_token], "丁单位")

    def test_detects_tobacco_alias_product_and_location_without_false_person(self):
        source = "山东中烟在文山产区生产品牌：文山雨露。"
        builder = MappingBuilder("domain123")
        anonymized = builder.anonymize(source)

        for sensitive_value in ("山东中烟", "文山", "文山雨露"):
            self.assertNotIn(sensitive_value, anonymized)
        self.assertEqual(builder.counts(), {"单位": 1, "产区": 1, "产品": 1})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    @override_settings(UIE_CATEGORY_THRESHOLDS={
        "person": 0.70,
        "organization": 0.55,
        "address": 0.60,
        "location": 0.60,
        "product": 0.60,
    })
    def test_model_conflict_resolution_filters_invalid_person_and_low_confidence(self):
        builder = MappingBuilder("conflict", ["person", "product", "organization", "location"])
        selected, rejected_count = _select_model_entities(builder, [
            {"text": "文山雨露", "category": "person", "probability": 0.99},
            {"text": "文山雨露", "category": "product", "probability": 0.82},
            {"text": "云南省", "category": "person", "probability": 0.99},
            {"text": "云南省", "category": "location", "probability": 0.80},
            {"text": "山东中烟", "category": "organization", "probability": 0.40},
        ])

        self.assertEqual(
            [(item["text"], item["category"]) for item in selected],
            [("文山雨露", "product"), ("云南省", "location")],
        )
        self.assertEqual(rejected_count, 3)

    def test_detects_pdf_style_spaces_inside_entities(self):
        source = "单 位：中 国 烟 草 总 公 司\n联 系 人：张 三\n138 0013 8000"
        builder = MappingBuilder("space123")
        anonymized = builder.anonymize(source)

        self.assertNotIn("中 国 烟 草 总 公 司", anonymized)
        self.assertNotIn("张 三", anonymized)
        self.assertNotIn("138 0013 8000", anonymized)
        self.assertEqual(builder.counts(), {"单位": 1, "人名": 1, "电话": 1})
        self.assertEqual(restore_text(anonymized, builder.export()), source)

    def test_replaces_all_normalized_variants_with_one_token_without_rewriting_sentences(self):
        builder = MappingBuilder(
            "consistent",
            ["organization"],
            [{"text": "陆良复烤厂", "category": "organization"}],
            token_namespace="LOCAL88",
        )
        source = "甲方为陆良复烤厂；复核仍由陆\u200b良 复 烤 厂负责，其他句子保持不变。"

        anonymized = builder.anonymize_registered(source)
        token = builder.original_to_token["陆良复烤厂"]

        self.assertEqual(token, "【LOCAL88-单001】")
        self.assertEqual(anonymized.count(token), 2)
        self.assertIn("；复核仍由", anonymized)
        self.assertIn("负责，其他句子保持不变。", anonymized)

    def test_reprocessing_keeps_active_tokens_and_old_files_restorable(self):
        first = MappingBuilder("stable", ["organization"], token_namespace="LOCAL88")
        first.register("甲单位", "organization")
        first.register("乙单位", "organization")
        first_output = first.anonymize_registered("甲单位与乙单位联合发布。")
        first_mapping = first.export()

        second = MappingBuilder(
            "stable",
            ["organization"],
            previous_mapping=first_mapping,
            token_namespace="LOCAL88",
        )
        second.register("乙单位", "organization")
        second.register("丙单位", "organization")
        second_mapping = second.export()

        self.assertEqual(
            second.original_to_token["乙单位"],
            first.original_to_token["乙单位"],
        )
        self.assertEqual(restore_text(first_output, second_mapping), "甲单位与乙单位联合发布。")
        self.assertNotEqual(
            second.original_to_token["丙单位"],
            first.original_to_token["甲单位"],
        )

    def test_does_not_treat_common_status_words_as_people(self):
        source = "审核状态：合格\n审批意见：同意\n项目\n方案\n申请\n费用\n单价\n说明\n文件\n安全\n管理"
        builder = MappingBuilder("safe1234", ["person"])
        self.assertEqual(builder.anonymize(source), source)
        self.assertEqual(builder.counts(), {})


class FileProcessorTests(TestCase):
    def setUp(self):
        self.directory = Path(tempfile.mkdtemp())
        self.source_text = "单位：中国烟草总公司\n联系人：张三\n电话：13800138000"

    def tearDown(self):
        shutil.rmtree(self.directory, ignore_errors=True)

    def _roundtrip(self, source, extension):
        builder = MappingBuilder("1234abcd")
        anonymized = self.directory / f"anonymous{extension}"
        restored = self.directory / f"restored{extension}"
        process_file(source, anonymized, builder.anonymize)
        self.assertTrue(anonymized.exists())
        process_file(anonymized, restored, lambda text: restore_text(text, builder.export()))
        self.assertTrue(restored.exists())
        return anonymized, restored

    def test_txt_roundtrip(self):
        source = self.directory / "sample.txt"
        source.write_text(self.source_text, encoding="utf-8")
        _, restored = self._roundtrip(source, ".txt")
        self.assertIn("中国烟草总公司", restored.read_text(encoding="utf-8"))

    def test_docx_roundtrip(self):
        source = self.directory / "sample.docx"
        document = Document()
        document.add_paragraph(self.source_text)
        document.save(source)
        anonymized, restored = self._roundtrip(source, ".docx")
        self.assertNotIn("张三", "\n".join(p.text for p in Document(anonymized).paragraphs))
        self.assertIn("张三", "\n".join(p.text for p in Document(restored).paragraphs))

    def test_docx_preserves_run_boundaries_and_formatting(self):
        source = self.directory / "styled.docx"
        document = Document()
        paragraph = document.add_paragraph()
        prefix = paragraph.add_run("委托单位：")
        prefix.bold = True
        sensitive_a = paragraph.add_run("山东")
        sensitive_a.italic = True
        sensitive_b = paragraph.add_run("中烟")
        sensitive_b.underline = True
        suffix = paragraph.add_run("，报告日期：2026年3月。")
        suffix.bold = True
        document.save(source)

        builder = MappingBuilder("styled")
        anonymized = self.directory / "styled-anonymized.docx"
        process_file(source, anonymized, builder.anonymize)
        result = Document(anonymized).paragraphs[0]

        self.assertEqual(len(result.runs), 4)
        self.assertEqual(result.text, "委托单位：【单001】，报告日期：2026年3月。")
        self.assertTrue(result.runs[0].bold)
        self.assertTrue(result.runs[1].italic)
        self.assertTrue(result.runs[2].underline)
        self.assertTrue(result.runs[3].bold)

    def test_docx_text_box_content_is_processed(self):
        source = self.directory / "textbox.docx"
        document = Document()
        text_box = parse_xml(
            f'<w:txbxContent {nsdecls("w")}>'
            '<w:p><w:r><w:t>项目负责人王建国，单位：中国烟草总公司</w:t></w:r></w:p>'
            '</w:txbxContent>'
        )
        document._element.body.append(text_box)
        document.save(source)

        anonymized, restored = self._roundtrip(source, ".docx")
        anonymized_text = "".join(
            node.text or "" for node in Document(anonymized)._element.xpath(".//w:txbxContent//w:t")
        )
        restored_text = "".join(
            node.text or "" for node in Document(restored)._element.xpath(".//w:txbxContent//w:t")
        )
        self.assertNotIn("王建国", anonymized_text)
        self.assertNotIn("中国烟草总公司", anonymized_text)
        self.assertIn("王建国", restored_text)
        self.assertIn("中国烟草总公司", restored_text)

    def test_docx_drawingml_text_from_converted_later_pages_is_processed(self):
        source = self.directory / "drawingml.docx"
        document = Document()
        paragraph = document.add_paragraph("前两页普通正文")
        drawing = parse_xml(
            f'<w:drawing {nsdecls("w", "a")}>'
            '<a:p><a:r><a:t>第三页单位：贵州烟叶复烤有限责任公司，与江苏中烟联合研究</a:t></a:r></a:p>'
            '</w:drawing>'
        )
        paragraph._p.append(drawing)
        document.save(source)

        anonymized, restored = self._roundtrip(source, ".docx")
        anonymized_text = "".join(
            node.text or "" for node in Document(anonymized)._element.xpath(".//a:t")
        )
        restored_text = "".join(
            node.text or "" for node in Document(restored)._element.xpath(".//a:t")
        )
        self.assertNotIn("贵州烟叶复烤有限责任公司", anonymized_text)
        self.assertNotIn("江苏中烟", anonymized_text)
        self.assertIn("贵州烟叶复烤有限责任公司", restored_text)
        self.assertIn("江苏中烟", restored_text)

    def test_xls_roundtrip(self):
        source = self.directory / "sample.xls"
        book = xlwt.Workbook()
        sheet = book.add_sheet("数据")
        sheet.write(0, 0, self.source_text)
        book.save(source)
        anonymized, restored = self._roundtrip(source, ".xls")
        self.assertNotIn("张三", xlrd.open_workbook(anonymized).sheet_by_index(0).cell_value(0, 0))
        self.assertIn("张三", xlrd.open_workbook(restored).sheet_by_index(0).cell_value(0, 0))

    def test_ofd_xml_roundtrip(self):
        source = self.directory / "sample.ofd"
        with zipfile.ZipFile(source, "w") as archive:
            archive.writestr("OFD.xml", f"<ofd:OFD xmlns:ofd='x'><ofd:Text>{self.source_text}</ofd:Text></ofd:OFD>")
        anonymized, restored = self._roundtrip(source, ".ofd")
        with zipfile.ZipFile(anonymized) as archive:
            self.assertNotIn("张三", archive.read("OFD.xml").decode())
        with zipfile.ZipFile(restored) as archive:
            self.assertIn("张三", archive.read("OFD.xml").decode())

    def test_ofd_split_text_code_nodes_are_processed(self):
        source = self.directory / "split-text.ofd"
        xml = (
            "<ofd:OFD xmlns:ofd='urn:ofd:test'><!-- vendor comment --><ofd:TextObject>"
            "<ofd:TextCode>联</ofd:TextCode><ofd:TextCode>系人：</ofd:TextCode>"
            "<ofd:TextCode>张</ofd:TextCode><ofd:TextCode>三</ofd:TextCode>"
            "</ofd:TextObject></ofd:OFD>"
        )
        with zipfile.ZipFile(source, "w") as archive:
            archive.writestr("OFD.xml", xml)

        anonymized, restored = self._roundtrip(source, ".ofd")
        with zipfile.ZipFile(anonymized) as archive:
            anonymized_xml = archive.read("OFD.xml").decode()
        with zipfile.ZipFile(restored) as archive:
            restored_xml = archive.read("OFD.xml").decode()
        self.assertNotIn("张三", anonymized_xml)
        self.assertIn("【人001】", anonymized_xml)
        self.assertIn("张三", restored_xml)

    def test_pdf_roundtrip(self):
        source = self.directory / "sample.pdf"
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        pdf = canvas.Canvas(str(source), pagesize=A4)
        pdf.setFont("STSong-Light", 12)
        pdf.drawString(50, 790, self.source_text.replace("\n", "  "))
        pdf.save()
        anonymized, restored = self._roundtrip(source, ".pdf")
        self.assertNotIn("张三", "".join(page.extract_text() or "" for page in __import__('pypdf').PdfReader(anonymized).pages))
        self.assertIn("张三", "".join(page.extract_text() or "" for page in __import__('pypdf').PdfReader(restored).pages))

    @override_settings(PDF_OCR_ENABLED=True)
    @patch("anonymizer.file_processors._ocr_pdf_pages")
    def test_scanned_pdf_uses_local_ocr_and_reports_progress(self, ocr_pages):
        source = self.directory / "scanned.pdf"
        pdf = canvas.Canvas(str(source), pagesize=A4)
        pdf.rect(50, 700, 300, 80)
        pdf.showPage()
        pdf.save()
        ocr_pages.return_value = {0: self.source_text}
        progress = []
        builder = MappingBuilder("scan123")
        anonymized = self.directory / "scanned-anonymized.pdf"

        process_file(
            source,
            anonymized,
            builder.anonymize,
            progress_callback=progress.append,
        )

        self.assertEqual(ocr_pages.call_args.args[:3], (source, [0], 1))
        extracted = "".join(page.extract_text() or "" for page in __import__('pypdf').PdfReader(anonymized).pages)
        self.assertNotIn("中国烟草总公司", extracted)
        self.assertNotIn("张三", extracted)
        self.assertTrue(any(item.get("stage") == "pdf_ocr" for item in progress))
        self.assertEqual(progress[-1]["stage"], "pdf_write")

    @override_settings(PDF_OCR_ENABLED=False)
    def test_scanned_pdf_fails_clearly_when_local_ocr_is_disabled(self):
        source = self.directory / "ocr-disabled.pdf"
        pdf = canvas.Canvas(str(source), pagesize=A4)
        pdf.rect(50, 700, 300, 80)
        pdf.save()

        with self.assertRaisesRegex(ProcessingError, "本地 OCR 已关闭"):
            process_file(source, self.directory / "unused.pdf", lambda text: text)


class TaskApiTests(TestCase):
    def setUp(self):
        self.media_directory = tempfile.mkdtemp()
        self.override = override_settings(
            MEDIA_ROOT=self.media_directory,
            MAPPING_ENCRYPTION_KEY="test-mapping-key",
            ANONYMIZATION_NAMESPACE="TESTHOST",
            REQUIRE_HUMAN_REVIEW=False,
        )
        self.override.enable()

    def tearDown(self):
        self.override.disable()
        shutil.rmtree(self.media_directory, ignore_errors=True)

    def test_txt_anonymize_and_restore_api(self):
        source = "单位：中国烟草总公司\n联系人：张三\n电话：13800138000"
        upload = SimpleUploadedFile("项目清单.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "person", "phone"]),
            "custom_entities": "",
        })
        self.assertEqual(response.status_code, 201, response.content)
        task = response.json()
        self.assertEqual(task["task_name"], "项目清单")
        self.assertEqual(task["stored_files"], {"original": True, "anonymized": True, "restore_input": False, "restored": False})
        download = self.client.get(task["anonymized_download_url"])
        anonymized = b"".join(download.streaming_content)
        self.assertNotIn("张三".encode(), anonymized)

        restore_upload = SimpleUploadedFile("AI处理稿.txt", anonymized + "\nAI补充内容".encode("utf-8"), content_type="text/plain")
        response = self.client.post(f"/api/tasks/{task['id']}/restore/", {"file": restore_upload})
        self.assertEqual(response.status_code, 200, response.content)
        self.assertEqual(response.json()["stored_files"], {"original": True, "anonymized": True, "restore_input": True, "restored": True})
        restored_download = self.client.get(response.json()["restored_download_url"])
        restored = b"".join(restored_download.streaming_content).decode("utf-8")
        self.assertIn("中国烟草总公司", restored)
        self.assertIn("张三", restored)
        self.assertIn("AI补充内容", restored)

    @override_settings(UIE_ENABLED=False)
    def test_docx_review_and_output_include_entities_after_third_page(self):
        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("第一页单位：中国烟草总公司。")
        document.add_page_break()
        document.add_paragraph("第二页为普通说明文字。")
        document.add_page_break()
        third_page = document.add_paragraph("第三页开头。")
        drawing = parse_xml(
            f'<w:drawing {nsdecls("w", "a")}>'
            '<a:p><a:r><a:t>遵义复烤厂中心实验室为贵州烟叶复烤有限责任公司中心实验室，'
            '与江苏中烟共计开展技术研究。</a:t></a:r></a:p>'
            '</w:drawing>'
        )
        third_page._p.append(drawing)
        document.add_page_break()
        document.add_paragraph("第四页：贵州毕节复烤厂持续推进技术改造。")
        document.save(buffer)
        upload = SimpleUploadedFile(
            "多页复烤报告.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        created = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization"]),
            "review_required": "true",
        })
        self.assertEqual(created.status_code, 201, created.content)
        task_id = created.json()["id"]
        review = self.client.get(f"/api/tasks/{task_id}/review/").json()
        preview_text = "\n".join(section["text"] for section in review["preview"])
        for expected in ("贵州烟叶复烤有限责任公司", "江苏中烟", "贵州毕节复烤厂"):
            self.assertIn(expected, preview_text)
        selected = [
            {"token": item["token"], "text": item["text"], "category": item["category"]}
            for item in review["entities"]
        ]

        confirmed = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {"selected_entities": selected},
            content_type="application/json",
        )
        self.assertEqual(confirmed.status_code, 200, confirmed.content)
        download = self.client.get(confirmed.json()["task"]["anonymized_download_url"])
        result_document = Document(io.BytesIO(b"".join(download.streaming_content)))
        output_text = "".join(
            node.text or "" for node in result_document._element.xpath(".//w:t | .//a:t")
        )
        for sensitive in (
            "中国烟草总公司",
            "遵义复烤厂",
            "贵州烟叶复烤有限责任公司",
            "江苏中烟",
            "贵州毕节复烤厂",
        ):
            self.assertNotIn(sensitive, output_text)

    @override_settings(MAX_UPLOAD_SIZE_MB=1)
    def test_rejects_oversized_upload_with_json_detail(self):
        upload = SimpleUploadedFile("过大.txt", b"a" * (1024 * 1024 + 1), content_type="text/plain")
        response = self.client.post("/api/tasks/", {"file": upload})

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "文件大小不能超过 1 MB。")

    @override_settings(MAX_UPLOAD_SIZE_MB=200)
    def test_stats_exposes_upload_limit_for_frontend_validation(self):
        response = self.client.get("/api/stats/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["max_upload_size_mb"], 200)

    @override_settings(UIE_ENABLED=False, PDF_OCR_ENABLED=True)
    @patch("anonymizer.file_processors._ocr_pdf_pages")
    def test_scanned_pdf_api_exposes_ocr_page_count_and_completed_progress(self, ocr_pages):
        ocr_pages.return_value = {0: "单位：中国烟草总公司\n联系人：张三"}
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        pdf.rect(50, 700, 300, 80)
        pdf.save()
        upload = SimpleUploadedFile("扫描公文.pdf", buffer.getvalue(), content_type="application/pdf")

        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "person"]),
        })

        self.assertEqual(response.status_code, 201, response.content)
        self.assertEqual(response.json()["ocr_page_count"], 1)
        self.assertEqual(response.json()["processing_progress"]["percent"], 100)
        self.assertEqual(response.json()["processing_progress"]["stage"], "completed")

    def test_api_returns_counts_for_realistic_unlabeled_content(self):
        source = (
            "云南中烟工业有限责任公司与红云红河烟草（集团）有限责任公司签订合同，"
            "项目负责人王建国，李娜负责复核，联系电话0871-63568888。"
        )
        upload = SimpleUploadedFile("真实公文语句.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "person", "phone"]),
            "custom_entities": "",
        })

        self.assertEqual(response.status_code, 201, response.content)
        task = response.json()
        self.assertEqual(task["entity_counts"], {"单位": 2, "人名": 2, "电话": 1})
        download = self.client.get(task["anonymized_download_url"])
        anonymized = b"".join(download.streaming_content).decode("utf-8")
        for sensitive_value in ("云南中烟工业有限责任公司", "红云红河烟草（集团）有限责任公司", "王建国", "李娜"):
            self.assertNotIn(sensitive_value, anonymized)

    @override_settings(UIE_ENABLED=True)
    @patch("anonymizer.views.predict_entities")
    def test_uie_model_entities_are_merged_and_mode_is_recorded(self, predict):
        predict.return_value = [{
            "text_index": 1,
            "text": "潘富昆",
            "category": "person",
            "start": 0,
            "end": 3,
            "probability": 0.99,
        }]
        upload = SimpleUploadedFile("潘富昆活动材料.txt", "活动按计划实施。".encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["person"]),
            "uie_mode": "resident",
        })

        self.assertEqual(response.status_code, 201, response.content)
        task = response.json()
        self.assertEqual(task["recognition_mode"], "resident")
        self.assertEqual(task["uie_detected_count"], 1)
        predict.assert_called_once()
        self.assertIn("潘富昆活动材料", predict.call_args.args[0])
        self.assertNotIn("潘富昆", task["display_name"])
        download = self.client.get(task["anonymized_download_url"])
        self.assertNotIn("潘富昆", download.headers["Content-Disposition"])

    def test_review_removes_false_positive_adds_missing_entity_and_reprocesses(self):
        source = "联系人：张三，内部代号星辰一号。"
        upload = SimpleUploadedFile("校正测试.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["person", "product"]),
            "custom_entities": "",
        })
        self.assertEqual(response.status_code, 201, response.content)
        task_id = response.json()["id"]

        review = self.client.get(f"/api/tasks/{task_id}/review/")
        self.assertEqual(review.status_code, 200, review.content)
        person_token = next(
            item["token"] for item in review.json()["entities"] if item["text"] == "张三"
        )
        response = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {"additions": "产品|星辰一号", "remove_tokens": [person_token]},
            content_type="application/json",
        )
        self.assertEqual(response.status_code, 200, response.content)
        self.assertEqual(response.json()["excluded_count"], 1)
        self.assertNotIn("张三", [item["text"] for item in response.json()["entities"]])
        self.assertIn("星辰一号", [item["text"] for item in response.json()["entities"]])

        task_data = response.json()["task"]
        download = self.client.get(task_data["anonymized_download_url"])
        anonymized = b"".join(download.streaming_content).decode("utf-8")
        self.assertIn("张三", anonymized)
        self.assertNotIn("星辰一号", anonymized)
        self.assertEqual(TrainingExample.objects.filter(action="rejected").count(), 1)

    def test_required_review_highlights_context_and_blocks_download_until_confirmation(self):
        source = "单位：中国烟草总公司。联系人：张三。内部产品代号星辰一号。"
        upload = SimpleUploadedFile("人工确认.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "person", "product"]),
            "review_required": "true",
        })

        self.assertEqual(response.status_code, 201, response.content)
        task = response.json()
        self.assertEqual(task["status"], "review")
        self.assertTrue(task["review_required"])
        self.assertFalse(task["review_confirmed"])
        self.assertIsNone(task["anonymized_download_url"])
        self.assertFalse(task["stored_files"]["anonymized"])
        blocked = self.client.get(f"/api/tasks/{task['id']}/download/anonymized/")
        self.assertEqual(blocked.status_code, 409)

        review = self.client.get(f"/api/tasks/{task['id']}/review/")
        self.assertEqual(review.status_code, 200, review.content)
        entities = review.json()["entities"]
        person = next(item for item in entities if item["text"] == "张三")
        self.assertEqual(person["occurrences"][0]["match"], "张三")
        self.assertIn("联系人", person["occurrences"][0]["prefix"])

        confirmed = self.client.post(
            f"/api/tasks/{task['id']}/review/",
            {
                "additions": "产品|星辰一号",
                "selected_entities": [{"token": person["token"], "category": "person"}],
            },
            content_type="application/json",
        )
        self.assertEqual(confirmed.status_code, 200, confirmed.content)
        completed = confirmed.json()["task"]
        self.assertEqual(completed["status"], "completed")
        self.assertTrue(completed["review_confirmed"])
        download = self.client.get(completed["anonymized_download_url"])
        anonymized = b"".join(download.streaming_content).decode("utf-8")
        self.assertIn("中国烟草总公司", anonymized)
        self.assertNotIn("张三", anonymized)
        self.assertNotIn("星辰一号", anonymized)

    @override_settings(UIE_ENABLED=False)
    def test_quality_stats_track_review_decisions_without_exposing_sensitive_text(self):
        source = "单位：中国烟草总公司。联系人：张三。内部密语火种。"
        upload = SimpleUploadedFile("质量指标.txt", source.encode("utf-8"), content_type="text/plain")
        created = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "person"]),
            "review_required": "true",
        })
        self.assertEqual(created.status_code, 201, created.content)
        task_id = created.json()["id"]
        review = self.client.get(f"/api/tasks/{task_id}/review/").json()
        person = next(item for item in review["entities"] if item["text"] == "张三")

        confirmed = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {
                "selected_entities": [{"token": person["token"], "text": "张三", "category": "person"}],
                "additions": "敏感项|密语火种",
            },
            content_type="application/json",
        )
        self.assertEqual(confirmed.status_code, 200, confirmed.content)

        response = self.client.get("/api/stats/")
        self.assertEqual(response.status_code, 200, response.content)
        payload = response.json()
        quality = payload["review_quality"]
        self.assertEqual(quality["reviewed_tasks"], 1)
        self.assertEqual(quality["selected_count"], 1)
        self.assertGreaterEqual(quality["rejected_count"], 1)
        self.assertEqual(quality["manual_added_count"], 1)
        self.assertGreater(payload["performance"]["measured_tasks"], 0)
        self.assertGreaterEqual(payload["entity_occurrences"], payload["entities"])
        self.assertNotIn("张三", response.content.decode("utf-8"))
        self.assertNotIn("密语火种", response.content.decode("utf-8"))

    def test_health_checks_database_readiness(self):
        response = self.client.get("/api/health/")
        self.assertEqual(response.status_code, 200, response.content)
        self.assertEqual(response.json()["database"], "ok")

    @override_settings(UIE_ENABLED=False)
    def test_full_preview_marks_all_occurrences_and_suggests_reviewed_alias_merge(self):
        source = "委托单位：陆良复烤厂。后续由陆良厂负责，陆良厂提交报告。"
        upload = SimpleUploadedFile("陆良复烤厂报告.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization"]),
            "review_required": "true",
        })
        self.assertEqual(response.status_code, 201, response.content)
        task_id = response.json()["id"]
        review = self.client.get(f"/api/tasks/{task_id}/review/").json()

        self.assertIn("陆良厂", [item["text"] for item in review["entities"]])
        group = next(item for item in review["alias_groups"] if "陆良厂" in item["members"])
        preview_spans = [span for section in review["preview"] for span in section["spans"]]
        self.assertEqual(sum(span["text"] == "陆良厂" for span in preview_spans), 2)

        selected = [
            {"token": item["token"], "text": item["text"], "category": item["category"]}
            for item in review["entities"]
        ]
        confirmed = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {
                "selected_entities": selected,
                "alias_groups": [dict(group, accepted=True, canonical="陆良复烤厂")],
            },
            content_type="application/json",
        )
        self.assertEqual(confirmed.status_code, 200, confirmed.content)
        download = self.client.get(confirmed.json()["task"]["anonymized_download_url"])
        anonymized = b"".join(download.streaming_content).decode("utf-8")
        self.assertEqual(anonymized.count("【TESTHOST-单001】"), 3)

    @override_settings(UIE_ENABLED=False)
    def test_review_saves_every_confirmed_category_and_keeps_old_output_restorable(self):
        source = "单位：中国烟草总公司。电话：13800138000。邮箱：local@example.com。密语火种。"
        upload = SimpleUploadedFile("一致性测试.txt", source.encode("utf-8"), content_type="text/plain")
        created = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["organization", "phone", "email"]),
            "review_required": "true",
        })
        self.assertEqual(created.status_code, 201, created.content)
        task_id = created.json()["id"]

        first_review = self.client.get(f"/api/tasks/{task_id}/review/").json()
        first_selected = [
            {"token": item["token"], "text": item["text"], "category": item["category"]}
            for item in first_review["entities"]
        ]
        first_confirmed = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {"selected_entities": first_selected},
            content_type="application/json",
        )
        self.assertEqual(first_confirmed.status_code, 200, first_confirmed.content)
        first_url = first_confirmed.json()["task"]["anonymized_download_url"]
        first_download = self.client.get(first_url)
        first_filename = Path(AnonymizationTask.objects.get(id=task_id).anonymized_file.name).name
        first_content = b"".join(first_download.streaming_content)

        saved_categories = set(RecognitionLabel.objects.values_list("category", flat=True))
        self.assertTrue({"organization", "phone", "email"}.issubset(saved_categories))
        for label in RecognitionLabel.objects.all():
            self.assertNotIn(decrypt_label(label), label.text_ciphertext)

        second_review = self.client.get(f"/api/tasks/{task_id}/review/").json()
        second_selected = [
            {"token": item["token"], "text": item["text"], "category": item["category"]}
            for item in second_review["entities"]
        ]
        corrected = self.client.post(
            f"/api/tasks/{task_id}/review/",
            {"selected_entities": second_selected, "additions": "敏感项|密语火种"},
            content_type="application/json",
        )
        self.assertEqual(corrected.status_code, 200, corrected.content)

        restore_upload = SimpleUploadedFile(first_filename, first_content, content_type="text/plain")
        restored = self.client.post(f"/api/tasks/{task_id}/restore/", {"file": restore_upload})
        self.assertEqual(restored.status_code, 200, restored.content)
        restored_download = self.client.get(restored.json()["restored_download_url"])
        self.assertEqual(b"".join(restored_download.streaming_content).decode("utf-8"), source)

    @override_settings(UIE_ENABLED=False)
    def test_training_document_machine_prelabel_human_save_and_jsonl_export(self):
        source = "单位：中国烟草总公司。联系人：张三。"
        upload = SimpleUploadedFile("训练样本.txt", source.encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/training/documents/", {"file": upload})
        self.assertEqual(response.status_code, 201, response.content)
        document = response.json()
        self.assertEqual(document["status"], "ready")
        self.assertTrue(document["preview"])
        selected = [
            {"text": item["text"], "category": item["category"]}
            for item in document["entities"]
        ]
        saved = self.client.post(
            f"/api/training/documents/{document['id']}/",
            {"selected_entities": selected, "additions": "单位|陆良厂"},
            content_type="application/json",
        )
        self.assertEqual(saved.status_code, 200, saved.content)
        self.assertEqual(saved.json()["status"], "labeled")
        self.assertTrue(RecognitionLabel.objects.filter(is_active=True).exists())
        self.assertTrue(TrainingExample.objects.filter(action="annotated").exists())
        exported = self.client.get("/api/training/export/")
        self.assertEqual(exported.status_code, 200)
        exported_text = b"".join(exported.streaming_content).decode("utf-8")
        self.assertIn("result_list", exported_text)
        self.assertIn('"prompt":"单位名称"', exported_text)
        self.assertIn('"prompt":"人名"', exported_text)
        self.assertIn('"result_list":[]', exported_text)
        self.assertTrue(TrainingDocument.objects.filter(id=document["id"]).exists())
        training_folder = Path(self.media_directory) / "training" / document["id"]
        self.assertTrue(training_folder.exists())
        deleted = self.client.delete(
            f"/api/training/documents/{document['id']}/",
            HTTP_X_TRAINING_DELETE_CONFIRM=document["id"],
        )
        self.assertEqual(deleted.status_code, 204)
        self.assertFalse(training_folder.exists())

    def test_training_labels_are_encrypted_versioned_and_used_by_later_tasks(self):
        response = self.client.post("/api/labels/", {"text": "李作英", "category": "person"}, content_type="application/json")
        self.assertEqual(response.status_code, 201, response.content)
        label = RecognitionLabel.objects.get(id=response.json()["id"])
        self.assertNotIn("李作英", label.text_ciphertext)
        self.assertEqual(decrypt_label(label), "李作英")
        self.assertEqual(TrainingExample.objects.count(), 1)

        response = self.client.patch(
            f"/api/labels/{label.id}/",
            {"text": "赵英桥", "category": "person"},
            content_type="application/json",
        )
        self.assertEqual(response.status_code, 200, response.content)
        self.assertEqual(TrainingExample.objects.count(), 2)
        sample = decrypt_mapping(TrainingExample.objects.first().payload_ciphertext)
        self.assertEqual(sample["after"]["text"], "赵英桥")

        upload = SimpleUploadedFile("后续任务.txt", "由赵英桥完成。".encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {
            "file": upload,
            "categories": json.dumps(["person"]),
            "uie_mode": "on_demand",
        })
        self.assertEqual(response.status_code, 201, response.content)
        download = self.client.get(response.json()["anonymized_download_url"])
        self.assertNotIn("赵英桥", b"".join(download.streaming_content).decode("utf-8"))

    @patch("anonymizer.views.set_runtime_mode")
    @patch("anonymizer.views.runtime_status")
    def test_model_runtime_control_api(self, status_mock, set_mode):
        status_mock.return_value = {"enabled": True, "available": True, "model": "uie-base", "resident_loaded": True}
        response = self.client.post("/api/model/runtime/", {"mode": "resident"}, content_type="application/json")
        self.assertEqual(response.status_code, 200, response.content)
        set_mode.assert_called_once_with("resident")
        self.assertTrue(response.json()["resident_loaded"])

    def test_anonymizes_download_filename_and_restores_formal_filename(self):
        source = "组长：潘富昆。\n成员：赵英桥、李作英。"
        upload = SimpleUploadedFile(
            "潘富昆技能竞赛活动方案.txt",
            source.encode("utf-8"),
            content_type="text/plain",
        )
        response = self.client.post(
            "/api/tasks/",
            {"file": upload, "categories": json.dumps(["person"]), "custom_entities": ""},
        )

        self.assertEqual(response.status_code, 201, response.content)
        task_data = response.json()
        task = AnonymizationTask.objects.get(id=task_data["id"])
        anonymized_name = Path(task.anonymized_file.name).name
        self.assertNotIn("潘富昆", anonymized_name)
        self.assertIn("ANON_TESTHOST-人", anonymized_name)
        self.assertEqual(task_data["display_name"], anonymized_name)
        self.assertNotIn("潘富昆", task_data["task_name"])

        download = self.client.get(task_data["anonymized_download_url"])
        anonymized_content = b"".join(download.streaming_content)
        self.assertNotIn("潘富昆", download.headers["Content-Disposition"])

        restore_upload = SimpleUploadedFile(anonymized_name, anonymized_content, content_type="text/plain")
        response = self.client.post(f"/api/tasks/{task.id}/restore/", {"file": restore_upload})
        self.assertEqual(response.status_code, 200, response.content)
        task.refresh_from_db()
        restored_name = Path(task.restored_file.name).name
        self.assertEqual(restored_name, "潘富昆技能竞赛活动方案_正式版.txt")
        self.assertEqual(response.json()["display_name"], restored_name)

    def test_delete_requires_confirmation_and_removes_files(self):
        upload = SimpleUploadedFile("采购方案_匿名.txt", "联系人：张三".encode("utf-8"), content_type="text/plain")
        response = self.client.post("/api/tasks/", {"file": upload, "categories": json.dumps(["person"])})
        self.assertEqual(response.status_code, 201, response.content)
        task_data = response.json()
        task = AnonymizationTask.objects.get(id=task_data["id"])
        task_directory = Path(self.media_directory) / "tasks" / str(task.id)
        self.assertEqual(task.task_name, "采购方案")
        self.assertTrue(task_directory.exists())

        response = self.client.delete(f"/api/tasks/{task.id}/")
        self.assertEqual(response.status_code, 400)
        self.assertTrue(AnonymizationTask.objects.filter(id=task.id).exists())

        response = self.client.delete(f"/api/tasks/{task.id}/", HTTP_X_TASK_DELETE_CONFIRM=str(task.id))
        self.assertEqual(response.status_code, 204)
        self.assertFalse(AnonymizationTask.objects.filter(id=task.id).exists())
        self.assertFalse(task_directory.exists())

    def test_rejects_spoofed_pdf(self):
        upload = SimpleUploadedFile("伪造.pdf", b"not really a PDF", content_type="application/pdf")
        response = self.client.post("/api/tasks/", {"file": upload})
        self.assertEqual(response.status_code, 400)
        self.assertIn("PDF", response.json()["detail"])

    def test_rejects_unsafe_ofd_archive_path(self):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("OFD.xml", "<OFD/>")
            archive.writestr("../escape.xml", "bad")
        upload = SimpleUploadedFile("恶意.ofd", buffer.getvalue(), content_type="application/octet-stream")
        response = self.client.post("/api/tasks/", {"file": upload})
        self.assertEqual(response.status_code, 400)
        self.assertIn("不安全", response.json()["detail"])


class PPStructureWorkerTests(TestCase):
    def test_builds_only_the_lite_ppstructure_modules(self):
        captured = {}
        created_models = []

        class FakeLayoutPipeline:
            def create_model(self, config, *args, **kwargs):
                created_models.append(config["model_name"])
                return config["model_name"]

        class FakeModule:
            @staticmethod
            def PPStructureV3(**kwargs):
                captured.update(kwargs)
                pipeline = FakeLayoutPipeline()
                captured["chart_model"] = pipeline.create_model(
                    {"model_name": "PP-Chart2Table"}
                )
                captured["layout_model"] = pipeline.create_model(
                    {"model_name": "PP-DocLayout-S"}
                )
                return object()

        with patch("anonymizer.ppstructure_worker._layout_pipeline_class", return_value=FakeLayoutPipeline), \
                patch.dict("sys.modules", {"paddleocr": FakeModule()}):
            build_ppstructure_pipeline()

        self.assertEqual(captured["layout_detection_model_name"], "PP-DocLayout-S")
        self.assertEqual(captured["text_detection_model_name"], "PP-OCRv5_mobile_det")
        self.assertEqual(captured["text_recognition_model_name"], "PP-OCRv5_mobile_rec")
        for option in (
            "use_doc_orientation_classify", "use_doc_unwarping", "use_textline_orientation",
            "use_seal_recognition", "use_table_recognition", "use_formula_recognition",
            "use_chart_recognition", "use_region_detection",
        ):
            self.assertFalse(captured[option], option)
        self.assertFalse(captured["enable_mkldnn"])
        self.assertNotIn("format_block_content", captured)
        self.assertIsNone(captured["chart_model"])
        self.assertEqual(captured["layout_model"], "PP-DocLayout-S")
        self.assertEqual(created_models, ["PP-DocLayout-S"])
        self.assertEqual(
            FakeLayoutPipeline().create_model({"model_name": "PP-Chart2Table"}),
            "PP-Chart2Table",
        )

    def test_prediction_uses_only_paddleocr_332_supported_options(self):
        captured = {}

        class FakePipeline:
            def predict(self, **kwargs):
                captured.update(kwargs)
                return []

        recognize_with_ppstructure(FakePipeline(), Path("check.png"))

        self.assertEqual(captured["input"], "check.png")
        self.assertNotIn("format_block_content", captured)

    def test_extracts_layout_order_and_keeps_text_outside_layout_blocks(self):
        class FakeResult:
            json = {
                "res": {
                    "parsing_res_list": [
                        {"block_label": "title", "block_content": "山东中烟"},
                        {"block_label": "text", "block_content": "联系人：潘富昆"},
                    ],
                    "overall_ocr_res": {
                        "rec_texts": ["山东中烟", "联系人：潘富昆", "表格内单位：云南中烟"]
                    },
                }
            }

        text = extract_ppstructure_text(FakeResult())

        self.assertEqual(text, "山东中烟\n联系人：潘富昆\n表格内单位：云南中烟")


class UIEWorkerTests(TestCase):
    def test_adds_guarded_legacy_aistudio_download_symbol(self):
        fake_package = ModuleType("aistudio_sdk")
        fake_hub = ModuleType("aistudio_sdk.hub")
        fake_package.hub = fake_hub

        with patch.dict(
            "sys.modules",
            {"aistudio_sdk": fake_package, "aistudio_sdk.hub": fake_hub},
        ):
            self.assertTrue(ensure_aistudio_download_compatibility())
            self.assertFalse(ensure_aistudio_download_compatibility())
            with self.assertRaisesRegex(RuntimeError, "已移除旧版 download API"):
                fake_hub.download(repo_id="unused")

    @override_settings(UIE_MANAGER_URL="file:///tmp/model.sock")
    def test_model_manager_rejects_non_loopback_url(self):
        with self.assertRaises(UIEProcessingError):
            _manager_url("/status")

    def test_worker_normalizes_numpy_like_probabilities_and_categories(self):
        class FakeEngine:
            def set_schema(self, schema):
                self.schema = schema

            def __call__(self, texts):
                return [{"人名": [{"text": "潘富昆", "start": 3, "end": 6, "probability": 0.98}]}]

        entities = _predict(FakeEngine(), {"texts": ["组长：潘富昆。"], "categories": ["person"]})
        self.assertEqual(entities, [{
            "text_index": 0,
            "text": "潘富昆",
            "category": "person",
            "start": 3,
            "end": 6,
            "probability": 0.98,
        }])
